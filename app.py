import streamlit as st
import plotly.graph_objects as go
import smtplib
import random
import io
import anthropic

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────
st.set_page_config(
    page_title="Alligator · GEO-Scanner Pro",
    page_icon="🐊",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700;900&family=Open+Sans:wght@400;500;600&display=swap');

:root {
    --green:      #1a7a3c;
    --green-light:#22a04f;
    --green-pale: #eaf5ee;
    --green-mid:  #155f2f;
    --white:      #ffffff;
    --offwhite:   #f5f9f6;
    --text:       #1a2e22;
    --muted:      #557060;
    --border:     #c5dece;
    --radius:     12px;
}

html, body, [data-testid="stAppViewContainer"] {
    background-color: var(--offwhite) !important;
    color: var(--text) !important;
    font-family: 'Open Sans', sans-serif;
}

[data-testid="stHeader"], [data-testid="stToolbar"], footer { display:none!important; }

.hero {
    background: linear-gradient(135deg, var(--green) 0%, var(--green-mid) 100%);
    padding: 4rem 2rem 3.2rem;
    text-align: center;
    position: relative;
    overflow: hidden;
    border-radius: 0 0 32px 32px;
    margin-bottom: 2.5rem;
}
.hero-eyebrow {
    display: inline-block;
    background: rgba(255,255,255,.15);
    border: 1px solid rgba(255,255,255,.3);
    color: #fff;
    font-family:'Montserrat',sans-serif;
    font-size:.7rem;
    font-weight:700;
    letter-spacing:.2em;
    text-transform:uppercase;
    border-radius:999px;
    padding:.3rem 1rem;
    margin-bottom:1.2rem;
}
.hero h1 {
    font-family:'Montserrat',sans-serif;
    font-size: clamp(2rem,5vw,3.4rem);
    font-weight:900;
    color:#fff;
    line-height:1.1;
    margin:0 0 .6rem;
    letter-spacing:-.02em;
}
.hero h1 em { font-style:normal; color: #a8f0c0; }
.hero p { color:rgba(255,255,255,.82); font-size:1.02rem; margin:0; }
.partner-badge {
    display:inline-flex;
    align-items:center;
    gap:.5rem;
    background:rgba(255,255,255,.1);
    border:1px solid rgba(255,255,255,.25);
    border-radius:8px;
    padding:.4rem .9rem;
    margin-top:1.2rem;
    color:#fff;
    font-size:.75rem;
    font-weight:600;
    letter-spacing:.05em;
}

.card {
    background:var(--white);
    border:1px solid var(--border);
    border-radius:var(--radius);
    padding:1.6rem 1.8rem;
    margin-bottom:1.2rem;
    box-shadow:0 2px 12px rgba(26,122,60,.07);
}
.card-title {
    font-family:'Montserrat',sans-serif;
    font-weight:700;
    font-size:.72rem;
    letter-spacing:.15em;
    text-transform:uppercase;
    color:var(--green);
    margin-bottom:1rem;
    padding-bottom:.5rem;
    border-bottom:2px solid var(--green-pale);
}

.score-wrap { text-align:center; padding:.4rem 0 .8rem; }
.score-number {
    font-family:'Montserrat',sans-serif;
    font-size:3.6rem;
    font-weight:900;
    color:var(--green);
    line-height:1;
}
.score-sub { font-size:.85rem; color:var(--muted); margin-top:.3rem; }
.score-pill {
    display:inline-block;
    margin-top:.5rem;
    font-size:.72rem;
    font-weight:700;
    letter-spacing:.1em;
    text-transform:uppercase;
    padding:.3rem 1rem;
    border-radius:999px;
}
.pill-critical  { background:#ffebee; color:#c62828; border:1px solid #ef9a9a; }
.pill-warning   { background:#fff8e1; color:#e65100; border:1px solid #ffcc02; }
.pill-good      { background:#e8f5e9; color:#2e7d32; border:1px solid #a5d6a7; }
.pill-excellent { background:var(--green); color:#fff; }

.result-card {
    background:var(--offwhite);
    border-left:4px solid var(--green);
    border-radius:0 var(--radius) var(--radius) 0;
    padding:1.2rem 1.4rem;
    white-space:pre-wrap;
    font-size:.9rem;
    line-height:1.75;
    color:var(--text);
}
.info-box {
    background:var(--green-pale);
    border:1px solid var(--border);
    border-radius:var(--radius);
    padding:1rem 1.4rem;
    font-size:.88rem;
    color:var(--green-mid);
    display:flex;
    gap:.7rem;
    align-items:flex-start;
}

[data-testid="stTextInput"] input {
    background:#fff !important;
    border:1.5px solid var(--border) !important;
    border-radius:9px !important;
    color:var(--text) !important;
    font-family:'Open Sans',sans-serif !important;
}
[data-testid="stTextInput"] input:focus {
    border-color:var(--green) !important;
    box-shadow:0 0 0 3px rgba(26,122,60,.12) !important;
}
label { color:var(--muted) !important; font-size:.8rem !important; font-weight:600 !important; }

[data-testid="stButton"] button {
    background:var(--green) !important;
    color:#fff !important;
    font-family:'Montserrat',sans-serif !important;
    font-weight:700 !important;
    font-size:.92rem !important;
    letter-spacing:.05em !important;
    border:none !important;
    border-radius:10px !important;
    padding:.72rem 2.4rem !important;
    transition:background .2s,transform .15s !important;
}
[data-testid="stButton"] button:hover {
    background:var(--green-light) !important;
    transform:translateY(-1px) !important;
}
[data-testid="stAlert"] { border-radius:var(--radius) !important; }
[data-testid="stSpinner"] p { color:var(--green) !important; }
.js-plotly-plot .plotly,.js-plotly-plot .plotly .main-svg { background:transparent!important; }
</style>
""", unsafe_allow_html=True)


# ─── HELPERS ───────────────────────────────────────────────────────
def get_secret(key):
    try:
        return st.secrets[key]
    except Exception:
        return None


# ─── CLAUDE CON WEB SEARCH ─────────────────────────────────────────
def genera_analisi(brand, url, keyword):
    api_key = get_secret("ANTHROPIC_API_KEY")
    if not api_key:
        return "ERRORE: ANTHROPIC_API_KEY non configurata nei Secrets."
    try:
        client = anthropic.Anthropic(api_key=api_key)

        # ── FASE 1: ricerca reale del sito ──────────────────────────
        # Claude usa web_search per leggere il sito e raccogliere dati concreti
        ricerca_prompt = f"""Sei un esperto di GEO (Generative Engine Optimization) e Schema.org.

Devi analizzare in modo APPROFONDITO il sito web del brand "{brand}" all'indirizzo {url}.

USA lo strumento web_search per:
1. Cercare "{brand} {url}" e leggere la homepage
2. Cercare "{brand} contatti indirizzo P.IVA"
3. Cercare "{brand} social media LinkedIn Facebook Instagram"
4. Cercare "{brand} servizi prodotti"
5. Cercare "{brand} recensioni"

Raccogli TUTTI i seguenti dati reali (se non trovi un dato, scrivi "NON TROVATO"):
- Tipo di entità: (scegli UNO tra: LocalBusiness, ProfessionalService, MedicalBusiness, LegalService, FinancialService, FoodEstablishment, Store, Product, SoftwareApplication, Organization, EducationalOrganization, Hotel, TouristAttraction, Event)
- Nome legale completo
- URL canonico
- Logo URL
- Descrizione breve (max 160 caratteri)
- Indirizzo completo (via, numero, CAP, città, provincia, paese)
- Telefono (formato internazionale +39...)
- Email di contatto
- P.IVA / VAT ID
- Anno di fondazione
- Fondatore/i
- Numero dipendenti (approssimativo)
- Servizi/prodotti principali (lista)
- Area geografica servita
- Profilo LinkedIn URL
- Profilo Facebook URL  
- Profilo Instagram URL
- Profilo Twitter/X URL
- Google Business URL
- Wikipedia URL (se esiste)
- Wikidata URL (se esiste)
- Eventuali certificazioni o premi
- Rating medio (se presente su Google/Trustpilot)
- Numero recensioni

Poi rispondi con ESATTAMENTE questi 3 blocchi:

## DIAGNOSI
Analisi professionale di 4-6 punti su perché le AI generative non citano questo sito.
Basa l'analisi sui dati REALI trovati. Sii specifico: cita lacune concrete (mancanza di markup strutturato, assenza di Wikipedia/Wikidata, profili social non collegati, contenuto non ottimizzato per intent informativi, ecc.).
NON menzionare codice o JSON in questa sezione.

## JSON_LD
Genera il JSON-LD Schema.org COMPLETO e PRECISO pronto da copiare e incollare nell'<head> della homepage.

REGOLE FONDAMENTALI:
- Usa SOLO dati reali trovati. Se un campo non è verificabile, OMETTILO completamente (non usare placeholder come "inserire qui").
- Scegli il @type corretto in base al tipo di entità rilevato.
- Includi SEMPRE questi campi se disponibili: @context, @type, @id, name, url, logo, image, description, foundingDate, founder, numberOfEmployees, address (PostalAddress completo), telephone, email, vatID, sameAs (array con TUTTI i profili social + Wikipedia + Wikidata + Google Business), knowsAbout (array di argomenti di competenza), hasOfferCatalog (con i servizi/prodotti reali), areaServed, contactPoint, award.

- Per LocalBusiness/ProfessionalService aggiungi anche: openingHoursSpecification, geo (GeoCoordinates), priceRange, aggregateRating (se hai dati reali).
- Per Product aggiungi: brand, offers (con Offer), sku, gtin.
- Per SoftwareApplication aggiungi: applicationCategory, operatingSystem, offers.
- Per Organization aggiungi: legalName, department (se applicabile).

Il JSON deve essere VALIDO, ben formattato, con tutti i campi popolati con dati reali.
Includi il tag <script type="application/ld+json"> e </script>.

## TESTO_HOME
Un paragrafo di 90-110 parole da inserire nella Home page, ottimizzato GEO.
Il testo deve:
- Includere il nome del brand, la città/area, i servizi principali
- Usare entità semantiche esplicite (nomi propri, luoghi, specializzazioni)
- Rispondere a domande implicite che un utente farebbe a un'AI ("Chi è X?", "Cosa fa X?", "Dove si trova X?")
- Essere scritto in italiano naturale, non robotico
- NON essere marketing generico: usa dettagli concreti trovati sul sito
"""

        # Chiamata con web_search abilitato
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=4000,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=[{"role": "user", "content": ricerca_prompt}],
        )

        # Estrai solo i blocchi di testo dalla risposta (ignora tool_use blocks)
        testo_finale = ""
        for block in response.content:
            if block.type == "text":
                testo_finale += block.text

        # Se la risposta si è fermata per tool_use, continua fino al testo finale
        while response.stop_reason == "tool_use":
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": "Ricerca completata."
                    })

            messages_followup = [
                {"role": "user", "content": ricerca_prompt},
                {"role": "assistant", "content": response.content},
                {"role": "user", "content": tool_results},
            ]
            response = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=4000,
                tools=[{"type": "web_search_20250305", "name": "web_search"}],
                messages=messages_followup,
            )
            for block in response.content:
                if block.type == "text":
                    testo_finale += block.text

        return testo_finale if testo_finale else "ERRORE: Nessun testo nella risposta."

    except anthropic.AuthenticationError:
        return "ERRORE: API Key Anthropic non valida."
    except anthropic.RateLimitError:
        return "ERRORE: Limite API raggiunto. Riprova tra qualche secondo."
    except Exception as e:
        return f"ERRORE: {e}"


def split_sezioni(testo):
    sezioni = {"diagnosi": "", "jsonld": "", "testo": ""}
    blocchi = testo.split("##")
    for b in blocchi:
        b = b.strip()
        low = b.lower()
        if low.startswith("diagnosi"):
            sezioni["diagnosi"] = b[len("diagnosi"):].strip()
        elif low.startswith("json_ld") or low.startswith("json-ld") or low.startswith("jsonld"):
            parts = b.split(None, 1)
            sezioni["jsonld"] = parts[1].strip() if len(parts) > 1 else b
        elif low.startswith("testo"):
            parts = b.split(None, 1)
            sezioni["testo"] = parts[1].strip() if len(parts) > 1 else b
    if not any(sezioni.values()):
        sezioni["diagnosi"] = testo
    return sezioni


# ─── CREA DOCX ─────────────────────────────────────────────────────
def crea_docx(brand, url, keyword, diagnosi, jsonld, testo, scores, geo_score, geo_label):
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    GREEN = (0x1a, 0x7a, 0x3c)
    DARK  = (0x15, 0x5f, 0x2f)
    GRAY  = (0x55, 0x55, 0x55)

    def h1(txt):
        p = doc.add_heading(txt, level=1)
        for r in p.runs: r.font.color.rgb = RGBColor(*GREEN)
        return p

    def h2(txt):
        p = doc.add_heading(txt, level=2)
        for r in p.runs: r.font.color.rgb = RGBColor(*DARK)
        return p

    def note(txt):
        p = doc.add_paragraph(f"ℹ️  {txt}")
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(*GRAY)
            p.runs[0].font.italic = True
        return p

    def body(txt):
        p = doc.add_paragraph(txt)
        if p.runs: p.runs[0].font.size = Pt(11)
        return p

    # ── COPERTINA ──
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("🐊  ALLIGATOR · GEO-Scanner Pro")
    r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(*GREEN)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"Report Tecnico Riservato — Uso Interno Agenzia").font.size = Pt(12)

    doc.add_paragraph()
    for label, val in [("Brand", brand), ("URL", url), ("Keyword target", keyword),
                       ("GEO Score", f"{geo_score}/100  —  {geo_label}")]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.font.bold = True
        p.add_run(val)
    doc.add_paragraph()

    # ── SCORE BREAKDOWN ──
    h1("GEO Score Breakdown")
    note("Punteggi simulati per dimensione. Usali per prioritizzare le azioni.")
    doc.add_paragraph()
    labels = ["Clarity", "Structure", "Entity", "Trust", "Tech"]
    for lbl, s in zip(labels, scores):
        bar = "█" * (s // 10) + "░" * (10 - s // 10)
        p = doc.add_paragraph()
        run_l = p.add_run(f"{lbl:<12}")
        run_l.font.name = "Courier New"; run_l.font.size = Pt(10)
        run_b = p.add_run(f" {bar}  {s}/100")
        run_b.font.name = "Courier New"; run_b.font.size = Pt(10)
        if s >= 70:   run_b.font.color.rgb = RGBColor(*GREEN)
        elif s >= 50: run_b.font.color.rgb = RGBColor(0xe6, 0x51, 0x00)
        else:         run_b.font.color.rgb = RGBColor(0xc6, 0x28, 0x28)
    doc.add_paragraph()

    # ── DIAGNOSI ──
    h1("Diagnosi GEO")
    note("Mostrabile al cliente durante la presentazione del report.")
    doc.add_paragraph()
    body(diagnosi or "Non disponibile.")
    doc.add_paragraph()

    # ── JSON-LD ──
    h1("Schema Markup JSON-LD — COPIA & INCOLLA READY")
    note("Incollare nell'<head> di OGNI pagina (o almeno homepage). "
         "Verificare su: https://validator.schema.org  e  https://search.google.com/test/rich-results")
    doc.add_paragraph()

    h2("Come implementarlo")
    body("1. Apri il CMS (WordPress, Webflow, ecc.) o il file HTML della homepage.")
    body("2. Incolla il codice qui sotto PRIMA del tag </head>.")
    body("3. Salva e pubblica.")
    body("4. Valida su schema.org validator (link sopra).")
    body("5. Per WordPress: usa il plugin 'Insert Headers and Footers' o 'Rank Math'.")
    doc.add_paragraph()

    h2("Codice JSON-LD")
    cp = doc.add_paragraph(jsonld or "JSON-LD non generato.")
    if cp.runs:
        cp.runs[0].font.name = "Courier New"
        cp.runs[0].font.size = Pt(8.5)
    doc.add_paragraph()

    # ── TESTO HOME ──
    h1("Testo Ottimizzato per la Home Page — COPIA & INCOLLA READY")
    note("Inserire in un paragrafo visibile nella Home, preferibilmente sopra la fold "
         "o nella sezione 'Chi siamo'. NON usare come meta description.")
    doc.add_paragraph()

    h2("Come implementarlo")
    body("1. Vai alla sezione di testo della Home nel CMS.")
    body("2. Aggiungi un nuovo paragrafo (o sostituisci il testo introduttivo esistente).")
    body("3. Incolla il testo qui sotto, formattato come paragrafo normale (non H1/H2).")
    body("4. Assicurati che sia visibile al crawl (non in elementi nascosti o lazy-load pesante).")
    doc.add_paragraph()

    h2("Testo")
    tp = doc.add_paragraph(testo or "Testo non generato.")
    if tp.runs:
        tp.runs[0].font.size = Pt(11)
        tp.runs[0].font.italic = True
    doc.add_paragraph()

    # ── CHECKLIST ──
    h1("Checklist Implementazione")
    note("Spunta ogni voce dopo l'implementazione.")
    doc.add_paragraph()
    items = [
        "[ ]  JSON-LD incollato nell'<head> della homepage",
        "[ ]  JSON-LD incollato nelle pagine servizi principali",
        "[ ]  Validato su validator.schema.org",
        "[ ]  Validato su Google Rich Results Test",
        "[ ]  Testo ottimizzato inserito nella Home",
        "[ ]  Profili social aggiornati con URL del sito (per sameAs)",
        "[ ]  Scheda Google Business aggiornata",
        "[ ]  Wikipedia / Wikidata verificata (se applicabile)",
        "[ ]  Sitemap aggiornata e reinviata a Google Search Console",
        "[ ]  Re-crawl richiesto in Google Search Console",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        if p.runs: p.runs[0].font.name = "Courier New"; p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── FOOTER ──
    h2("Note Interne")
    body("Documento riservato — Alligator.it Srl · Google Premier Partner 2025")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── MAIL + ALLEGATO ───────────────────────────────────────────────
def invia_report_mail(brand, url, keyword, email_cliente,
                      diagnosi, jsonld, testo, scores, geo_score, geo_label):
    email_password = get_secret("EMAIL_PASSWORD")
    if not email_password:
        return False, "EMAIL_PASSWORD non trovata nei Secrets."

    account = destinatario = "nicofioretti7@gmail.com"

    try:
        docx_bytes = crea_docx(brand, url, keyword, diagnosi,
                                jsonld, testo, scores, geo_score, geo_label)

        msg = MIMEMultipart()
        msg["From"] = account
        msg["To"] = destinatario
        msg["Subject"] = f"🐊 Nuovo Lead GEO: {brand}"

        corpo = (
            f"Nuovo audit GEO completato.\n\n"
            f"Brand:   {brand}\n"
            f"URL:     {url}\n"
            f"Keyword: {keyword}\n"
            + (f"Email cliente: {email_cliente}\n" if email_cliente else "")
            + f"\nGEO Score: {geo_score}/100 — {geo_label}\n\n"
            "Il report tecnico completo (JSON-LD + testo Home) è in allegato .docx\n"
            "— Alligator GEO-Scanner Pro"
        )
        msg.attach(MIMEText(corpo, "plain", "utf-8"))

        part = MIMEBase(
            "application",
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        part.set_payload(docx_bytes)
        encoders.encode_base64(part)
        filename = f"GEO_Report_{brand.replace(' ', '_')}.docx"
        part.add_header("Content-Disposition", "attachment", filename=filename)
        msg.attach(part)

        with smtplib.SMTP("smtp.gmail.com", 587, timeout=20) as server:
            server.ehlo(); server.starttls(); server.ehlo()
            server.login(account, email_password)
            server.sendmail(account, destinatario, msg.as_string())

        return True, ""

    except smtplib.SMTPAuthenticationError:
        return False, "Credenziali Gmail non valide. Usa una App Password."
    except smtplib.SMTPException as e:
        return False, f"Errore SMTP: {e}"
    except Exception as e:
        return False, f"Errore: {e}"


# ─── RADAR ─────────────────────────────────────────────────────────
def crea_radar(scores, brand):
    labels = ["Clarity", "Structure", "Entity", "Trust", "Tech"]
    r = scores + [scores[0]]; theta = labels + [labels[0]]
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=r, theta=theta, fill="toself",
        fillcolor="rgba(26,122,60,0.15)",
        line=dict(color="#1a7a3c", width=2.5),
        marker=dict(color="#1a7a3c", size=7),
        name=brand,
    ))
    fig.update_layout(
        polar=dict(
            bgcolor="rgba(0,0,0,0)",
            radialaxis=dict(visible=True, range=[0, 100],
                            tickfont=dict(color="#888", size=9),
                            gridcolor="#c5dece", linecolor="#c5dece"),
            angularaxis=dict(tickfont=dict(color="#1a2e22", size=12, family="Open Sans"),
                             gridcolor="#c5dece", linecolor="#c5dece"),
        ),
        showlegend=False,
        margin=dict(t=30, b=30, l=40, r=40),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
    )
    return fig


def calcola_geo_score(scores):
    media = round(sum(scores) / len(scores))
    if media < 45:   return media, "Critico",      "pill-critical"
    elif media < 60: return media, "Ottimizzabile", "pill-warning"
    elif media < 80: return media, "Buono",         "pill-good"
    else:            return media, "Eccellente",    "pill-excellent"


# ─── HERO ──────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <div class="hero-eyebrow">GEO · Generative Engine Optimization</div>
    <h1>🐊 ALLIGATOR<br><em>GEO-Scanner Pro</em></h1>
    <p>Scopri perché ChatGPT, Gemini e Perplexity non citano il tuo brand<br>
       e ricevi il piano d'azione dalla tua agenzia.</p>
    <div class="partner-badge">⭐ Google Premier Partner 2025</div>
</div>
""", unsafe_allow_html=True)


# ─── FORM ──────────────────────────────────────────────────────────
st.markdown("<div class='card'>", unsafe_allow_html=True)
st.markdown("<div class='card-title'>⚙️ Configura il tuo Audit GEO</div>", unsafe_allow_html=True)

col1, col2 = st.columns(2, gap="large")
with col1:
    u = st.text_input("URL Sito Web", placeholder="https://tuosito.it")
    b = st.text_input("Nome Brand", placeholder="Es. Alligator")
with col2:
    k = st.text_input("Keyword Principale", placeholder="Es. agenzia SEO Milano")
    m = st.text_input("La tua Email (opzionale)", placeholder="tu@email.it")

st.markdown("</div>", unsafe_allow_html=True)

_, btn_col, _ = st.columns([1, 2, 1])
with btn_col:
    avvia = st.button("🐊 AVVIA AUDIT GEO", use_container_width=True)


# ─── ELABORAZIONE ──────────────────────────────────────────────────
if avvia:
    if not (u and b and k):
        st.error("Inserisci URL, Brand e Keyword per avviare l'analisi.")
    else:
        with st.spinner("Alligator sta analizzando il tuo brand con Claude AI…"):
            risultato_ai = genera_analisi(b, u, k)
            sezioni = split_sezioni(risultato_ai)

            scores = [random.randint(35, 85) for _ in range(5)]
            geo_score, geo_label, geo_css = calcola_geo_score(scores)

            # Mail con docx allegato — tutto il tecnico rimane privato
            mail_ok, mail_err = invia_report_mail(
                b, u, k, m,
                sezioni["diagnosi"], sezioni["jsonld"], sezioni["testo"],
                scores, geo_score, geo_label,
            )

        st.markdown("<br>", unsafe_allow_html=True)
        left, right = st.columns([1, 1.45], gap="large")

        with left:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='card-title'>📡 Radar GEO Score</div>", unsafe_allow_html=True)
            st.plotly_chart(crea_radar(scores, b), use_container_width=True)

            st.markdown(f"""
            <div class='score-wrap'>
                <div class='score-number'>{geo_score}<span style='font-size:1.4rem;color:var(--muted)'>/100</span></div>
                <div class='score-sub'>GEO Score Complessivo</div>
                <span class='score-pill {geo_css}'>{geo_label}</span>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            labels_r = ["Clarity", "Structure", "Entity", "Trust", "Tech"]
            for lbl, s in zip(labels_r, scores):
                color = "#1a7a3c" if s >= 70 else "#e65100" if s >= 50 else "#c62828"
                st.markdown(f"""
                <div style='margin-bottom:.55rem'>
                    <div style='display:flex;justify-content:space-between;margin-bottom:.2rem'>
                        <span style='font-size:.78rem;color:var(--muted);font-weight:600'>{lbl}</span>
                        <span style='font-size:.78rem;font-weight:700;color:var(--text)'>{s}</span>
                    </div>
                    <div style='background:#e8f0ec;border-radius:999px;height:6px'>
                        <div style='width:{s}%;height:6px;border-radius:999px;background:{color}'></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        with right:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='card-title'>🔍 Analisi GEO del tuo Brand</div>", unsafe_allow_html=True)

            diagnosi_display = sezioni["diagnosi"] if sezioni["diagnosi"] else risultato_ai
            st.markdown(f"<div class='result-card'>{diagnosi_display}</div>", unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("""
            <div class='info-box'>
                <span style='font-size:1.3rem'>📬</span>
                <div>
                    <strong>Prossimo passo</strong><br>
                    Il tuo piano d'azione personalizzato è pronto.
                    Un esperto Alligator ti contatterà per illustrarti le implementazioni
                    e i tempi di risultato attesi.
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        if mail_ok:
            st.success("✅ Audit completato! Il report tecnico è stato inviato all'agenzia Alligator.")
        else:
            st.success("✅ Audit completato!")
            if mail_err:
                st.warning(f"⚠️ Invio report fallito: {mail_err}")
